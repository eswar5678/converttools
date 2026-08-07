#!/usr/bin/env python3
"""Generate sample files for testing ConvertTools."""
import math
import os
import zipfile

from docx import Document
from docx.shared import Pt, Inches
from PIL import Image, ImageDraw

BASE = os.path.dirname(os.path.abspath(__file__))


def doc(x): return os.path.join(BASE, x)


def make_docx_report():
    d = Document()
    d.add_heading('Quarterly Report', level=0)
    d.add_heading('1. Overview', level=1)
    p = d.add_paragraph('Revenue grew ')
    p.add_run('12%').bold = True
    p.add_run(' this quarter, driven by ')
    p.add_run('strong').italic = True
    p.add_run(' subscription sales across all regions.')
    d.add_heading('2. Key Numbers', level=1)
    t = d.add_table(rows=3, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = 'Metric'
    t.rows[0].cells[1].text = 'Value'
    t.rows[1].cells[0].text = 'Revenue'
    t.rows[1].cells[1].text = '$1.2M'
    t.rows[2].cells[0].text = 'New users'
    t.rows[2].cells[1].text = '8,450'
    d.add_heading('3. Next Steps', level=1)
    for item in ('Expand to two new markets', 'Hire three engineers', 'Ship the mobile app'):
        d.add_paragraph(item, style='List Bullet')
    d.save(doc('report.docx'))
    print('report.docx')


def make_docx_letter():
    d = Document()
    d.add_heading('Cover Letter', level=0)
    d.add_paragraph('Dear Hiring Team,')
    d.add_paragraph(
        'I am writing to express my interest in the Web Developer role. '
        'I have five years of experience building fast, accessible interfaces '
        'and I care deeply about user privacy.'
    )
    d.add_paragraph('Kind regards,')
    d.add_paragraph('A. Candidate')
    d.save(doc('letter.docx'))
    print('letter.docx')


def gradient_image(path, c1, c2, size=(1200, 900), fmt='PNG'):
    img = Image.new('RGB', size)
    draw = ImageDraw.Draw(img)
    w, h = size
    for y in range(h):
        t = y / h
        col = tuple(int(c1[i] + (c2[i] - c1[i]) * t) for i in range(3))
        draw.line([(0, y), (w, y)], fill=col)
    draw.ellipse([w * 0.2, h * 0.25, w * 0.55, h * 0.75], fill=(255, 255, 255))
    draw.rectangle([w * 0.55, h * 0.35, w * 0.85, h * 0.7], fill=(30, 27, 58))
    if fmt == 'JPEG':
        img.save(path, 'JPEG', quality=90)
    else:
        img.save(path, 'PNG')
    print(path)


def make_plain_pdf(path, title, body):
    """Hand-craft a minimal valid one-page PDF with text."""
    stream = f"BT /F1 22 Tf 56 780 Td ({title}) Tj ET\n"
    y = 730
    for line in body:
        safe = line.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')
        stream += f"BT /F1 12 Tf 56 {y} Td ({safe}) Tj ET\n"
        y -= 22
    stream_b = stream.encode('latin-1')
    objs = []
    objs.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objs.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objs.append(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>")
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    objs.append(b"<< /Length %d >>\nstream\n" % len(stream_b) + stream_b + b"\nendstream")
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, o in enumerate(objs, 1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + o + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objs)+1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += ("%010d 00000 n \n" % off).encode('latin-1')
    out += (b"trailer\n<< /Size " + str(len(objs) + 1).encode() + b" /Root 1 0 R >>\nstartxref\n"
            + str(xref_pos).encode() + b"\n%%EOF")
    with open(path, 'wb') as f:
        f.write(bytes(out))
    print(path)


if __name__ == '__main__':
    make_docx_report()
    make_docx_letter()
    gradient_image(doc('photo-1.png'), (139, 92, 246), (236, 72, 153), fmt='PNG')
    gradient_image(doc('photo-2.jpg'), (245, 158, 11), (34, 197, 94), fmt='JPEG')
    with open(doc('story.txt'), 'w') as f:
        f.write("The Little Tool That Could\n\n"
                "Once upon a time, a small static website decided it would never upload anyone's files. "
                "It learned to merge PDFs, to turn photos into documents, and to pull text out of PDFs, "
                "all inside the visitor's own browser.\n\n"
                "People asked: how can this be free? The little site replied: because there is no server "
                "bill when your processor does the work. And everyone downloaded their converted files "
                "happily ever after.\n\nThe end.\n")
    print('story.txt')
    make_plain_pdf(doc('chapter-1.pdf'), 'Chapter One', [
        'It was a bright cold day in April,',
        'and the clocks were striking thirteen.',
        'Every document on the desk waited patiently',
        'to be merged into a single tidy PDF.',
    ])
    make_plain_pdf(doc('chapter-2.pdf'), 'Chapter Two', [
        'The second chapter arrives with its own page,',
        'bringing more words and a hopeful ending.',
        'When these two files are combined,',
        'the story becomes complete.',
    ])
    print('--- all fixtures created ---')
